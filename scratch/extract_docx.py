import os
try:
    import docx
except ImportError:
    import pip
    pip.main(['install', 'python-docx'])
    import docx

def extract_text(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return '\n'.join(full_text)

if __name__ == "__main__":
    docx_path = r'c:\Users\egas1\Downloads\Stitch_\catalog.docx'
    if os.path.exists(docx_path):
        text = extract_text(docx_path)
        with open(r'c:\Users\egas1\Downloads\Stitch_\catalog_extracted.txt', 'w', encoding='utf-8') as f:
            f.write(text)
        print("Extraction complete.")
    else:
        print("File not found.")
